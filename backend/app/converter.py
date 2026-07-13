"""
Input converter — normalizes DOCX, Markdown, LaTeX, and plain text
into the internal document representation used by format plugins.
"""

import re
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

import structlog
from docx import Document as DocxDocument

log = structlog.get_logger()

# Supplementary section heading patterns
_SUPPLEMENTARY_RE = re.compile(
    r"^(supplementary|supplemental|supporting\s+information|appendix)",
    re.IGNORECASE,
)


def convert_input(
    input_path: str,
    bib_path: Optional[str] = None,
    figures_dir: Optional[str] = None,
) -> dict:
    """
    Convert any supported input format to internal document representation.
    Returns: { "items": [...], "ris_data": [...] | None, "metadata": {...} }

    figures_dir: directory where extracted/uploaded images are stored.
                 Converter will save DOCX-embedded images here and record
                 image items referencing relative paths within this dir.
    """
    path = Path(input_path)
    suffix = path.suffix.lower()
    figures_path = Path(figures_dir) if figures_dir else None

    if suffix == ".docx":
        items = _read_docx(input_path, figures_path)
    elif suffix in (".md", ".markdown"):
        items = _read_markdown(input_path, figures_path)
    elif suffix == ".tex":
        items = _read_latex(input_path)
    elif suffix == ".txt":
        items = _read_plaintext(input_path)
    else:
        raise ValueError(f"Unsupported input format: {suffix}")

    # Parse bibliography if provided
    ris_data = None
    if bib_path:
        bib_path_obj = Path(bib_path)
        if bib_path_obj.exists():
            if bib_path_obj.suffix.lower() == ".ris":
                ris_data = _parse_ris(bib_path)
            elif bib_path_obj.suffix.lower() == ".bib":
                ris_data = _parse_bibtex(bib_path)

    return {"items": items, "ris_data": ris_data, "metadata": _extract_metadata(items)}


# ---------------------------------------------------------------------------
# DOCX reader
# ---------------------------------------------------------------------------

def _read_docx(path: str, figures_path: Optional[Path] = None) -> list:
    """Read DOCX and classify each element into the internal item schema.

    Extracts embedded images from inline shapes and inserts them as
    {"type": "image", ...} items at the correct position.
    """
    doc = DocxDocument(path)
    items = []
    found_abstract = False
    found_keywords = False
    image_counter = [0]  # mutable for nested helper

    # Build a map: paragraph element XML id → list of image items
    # so we can insert images after the paragraph that contains them.
    para_image_map: dict[int, list] = {}

    if figures_path:
        figures_path.mkdir(parents=True, exist_ok=True)
        para_image_map = _extract_docx_images(doc, figures_path, image_counter)

    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style_name = para.style.name if para.style else ""
        runs = _extract_runs(para)

        if text:
            # Classify by style and content
            if style_name == "Title" or (not items and not any(
                s in style_name for s in ["Heading", "Normal"]
            )):
                item_type = "title"
            elif "Heading 1" in style_name:
                item_type = "heading1"
            elif "Heading 2" in style_name:
                item_type = "heading2"
            elif "Heading 3" in style_name:
                item_type = "heading3"
            elif text.lower().startswith("abstract"):
                item_type = "abstract_heading"
                found_abstract = True
            elif found_abstract and not found_keywords and not any(
                h in style_name for h in ["Heading"]
            ):
                item_type = "abstract"
            elif text.lower().startswith("keywords"):
                item_type = "keywords"
                found_keywords = True
            elif re.match(r"^\[?\d+\]?\s+\w", text) and len(text) > 30:
                item_type = "reference"
            elif re.match(r"^(Table|Figure)\s+\d+", text, re.IGNORECASE):
                item_type = "table_caption" if text.lower().startswith("table") else "figure_caption"
            elif _SUPPLEMENTARY_RE.match(text):
                item_type = "supplementary_heading"
            else:
                item_type = "paragraph"

            items.append({"type": item_type, "text": text, "runs": runs})

        # Insert any images that were embedded in this paragraph
        for img_item in para_image_map.get(para_idx, []):
            items.append(img_item)

    # Extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append({
                    "text": cell.text.strip(),
                    "gridspan": 1,
                    "runs": [{"text": cell.text.strip(), "bold": False, "italic": False,
                               "superscript": False, "subscript": False}],
                    "vmerge_continue": False,
                })
            rows.append(cells)
        items.append({"type": "table", "text": "", "runs": [], "rows": rows})

    return items


def _extract_docx_images(doc, figures_path: Path, counter: list) -> dict:
    """Extract embedded images from a DOCX document.

    Returns a dict mapping paragraph index → list of image items.
    Images are saved to figures_path as fig_001.png, fig_002.png, etc.
    """
    from lxml import etree

    # Namespace map used in DOCX XML
    nsmap = {
        "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "r":   "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }

    para_image_map: dict[int, list] = {}

    for para_idx, para in enumerate(doc.paragraphs):
        para_elem = para._element
        # Find all blip elements (image references) in this paragraph
        blips = para_elem.findall(
            ".//a:blip",
            namespaces={"a": "http://schemas.openxmlformats.org/drawingml/2006/main"},
        )
        for blip in blips:
            r_embed = blip.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not r_embed:
                continue
            try:
                image_part = doc.part.related_parts.get(r_embed)
                if image_part is None:
                    continue
                # Determine extension from content type
                ct = image_part.content_type  # e.g. "image/png"
                ext_map = {
                    "image/png": "png",
                    "image/jpeg": "jpg",
                    "image/gif": "gif",
                    "image/svg+xml": "svg",
                    "image/tiff": "tif",
                    "image/bmp": "bmp",
                    "image/webp": "webp",
                    "image/emf": "emf",
                    "image/wmf": "wmf",
                }
                ext = ext_map.get(ct, "png")
                # Skip EMF/WMF (Windows metafiles) — not embeddable in PDF
                if ext in ("emf", "wmf"):
                    continue

                counter[0] += 1
                fname = f"fig_{counter[0]:03d}.{ext}"
                dest = figures_path / fname
                dest.write_bytes(image_part.blob)

                img_item = {
                    "type": "image",
                    "path": str(dest),
                    "alt": f"Figure {counter[0]}",
                    "caption": "",
                }
                para_image_map.setdefault(para_idx, []).append(img_item)
                log.info("docx_image_extracted", fname=fname, size=len(image_part.blob))
            except Exception as e:
                log.warning("docx_image_extract_failed", error=str(e))

    return para_image_map


# ---------------------------------------------------------------------------
# Markdown reader
# ---------------------------------------------------------------------------

def _read_markdown(path: str, figures_path: Optional[Path] = None) -> list:
    """Convert Markdown to internal items via Pandoc JSON AST."""
    raw_text = Path(path).read_text(encoding="utf-8")

    # Pre-scan for image references before Pandoc (handles local file paths)
    image_refs = _scan_markdown_images(raw_text, figures_path)

    try:
        result = subprocess.run(
            ["pandoc", path, "-t", "json"],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            items = _pandoc_json_to_items(result.stdout, figures_path)
            # Inject image items from pre-scan if Pandoc didn't capture them
            return items
    except (subprocess.TimeoutExpired, FileNotFoundError):
        pass

    # Fallback: simple regex-based Markdown parser
    return _parse_markdown_simple(raw_text, figures_path)


def _scan_markdown_images(text: str, figures_path: Optional[Path]) -> list:
    """Scan Markdown text for ![alt](path) image references.

    Returns list of image items. If figures_path is given and the path
    is a local file that exists there, records the absolute path.
    """
    image_items = []
    pattern = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
    for m in pattern.finditer(text):
        alt = m.group(1)
        src = m.group(2).strip()
        # Resolve local path
        resolved = src
        if figures_path and not src.startswith(("http://", "https://", "data:")):
            candidate = figures_path / Path(src).name
            if candidate.exists():
                resolved = str(candidate)
        image_items.append({
            "type": "image",
            "path": resolved,
            "alt": alt or "Figure",
            "caption": alt or "",
        })
    return image_items


# ---------------------------------------------------------------------------
# LaTeX / plain text readers
# ---------------------------------------------------------------------------

def _read_latex(path: str) -> list:
    """Convert LaTeX to internal items via Pandoc."""
    try:
        result = subprocess.run(
            ["pandoc", path, "-f", "latex", "-t", "json"],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return _pandoc_json_to_items(result.stdout)
    except (subprocess.TimeoutExpired, FileNotFoundError):
        pass

    return _parse_markdown_simple(Path(path).read_text(encoding="utf-8"))


def _read_plaintext(path: str) -> list:
    """Parse plain text into items using heuristics."""
    text = Path(path).read_text(encoding="utf-8")
    return _parse_markdown_simple(text)


# ---------------------------------------------------------------------------
# Simple Markdown parser (fallback)
# ---------------------------------------------------------------------------

def _parse_markdown_simple(text: str, figures_path: Optional[Path] = None) -> list:
    """Simple Markdown/plain text parser — fallback when Pandoc unavailable."""
    items = []
    lines = text.split("\n")
    i = 0
    in_abstract = False

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Image syntax: ![alt](path)
        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if img_match:
            alt = img_match.group(1)
            src = img_match.group(2).strip()
            resolved = src
            if figures_path and not src.startswith(("http://", "https://", "data:")):
                candidate = figures_path / Path(src).name
                if candidate.exists():
                    resolved = str(candidate)
            items.append({
                "type": "image",
                "path": resolved,
                "alt": alt or "Figure",
                "caption": alt or "",
            })
            i += 1
            continue

        # ATX headings
        if line.startswith("# "):
            items.append({"type": "heading1", "text": line[2:].strip(),
                          "runs": [{"text": line[2:].strip(), "bold": True, "italic": False,
                                    "superscript": False, "subscript": False}]})
        elif line.startswith("## "):
            heading_text = line[3:].strip()
            tl = heading_text.lower().rstrip(".")
            if tl == "abstract":
                items.append({"type": "abstract_heading", "text": heading_text,
                              "runs": [{"text": heading_text, "bold": True, "italic": False,
                                        "superscript": False, "subscript": False}]})
                in_abstract = True
            elif _SUPPLEMENTARY_RE.match(heading_text):
                items.append({"type": "supplementary_heading", "text": heading_text,
                              "runs": [{"text": heading_text, "bold": True, "italic": False,
                                        "superscript": False, "subscript": False}]})
                in_abstract = False
            else:
                items.append({"type": "heading2", "text": heading_text,
                              "runs": [{"text": heading_text, "bold": True, "italic": False,
                                        "superscript": False, "subscript": False}]})
        elif line.startswith("### "):
            items.append({"type": "heading3", "text": line[4:].strip(),
                          "runs": [{"text": line[4:].strip(), "bold": True, "italic": False,
                                    "superscript": False, "subscript": False}]})
        elif line.lower().startswith("abstract"):
            in_abstract = True
            items.append({"type": "abstract_heading", "text": line,
                          "runs": [{"text": line, "bold": True, "italic": False,
                                    "superscript": False, "subscript": False}]})
        elif line.lower().startswith("keywords"):
            in_abstract = False
            items.append({"type": "keywords", "text": line,
                          "runs": [{"text": line, "bold": False, "italic": False,
                                    "superscript": False, "subscript": False}]})
        elif re.match(r"^\[?\d+\]?\s+\w", line) and len(line) > 30:
            items.append({"type": "reference", "text": line,
                          "runs": [{"text": line, "bold": False, "italic": False,
                                    "superscript": False, "subscript": False}]})
        elif re.match(r"^(Table|Figure)\s+\d+", line, re.IGNORECASE):
            t = "table_caption" if line.lower().startswith("table") else "figure_caption"
            items.append({"type": t, "text": line,
                          "runs": [{"text": line, "bold": False, "italic": True,
                                    "superscript": False, "subscript": False}]})
        else:
            item_type = "abstract" if in_abstract else "paragraph"
            if not items and len(line) < 150 and not line.endswith("."):
                item_type = "title"
                in_abstract = False
            runs = _parse_inline_markdown(line)
            items.append({"type": item_type, "text": line, "runs": runs})

        i += 1

    return items


def _parse_inline_markdown(text: str) -> list:
    """Parse inline Markdown formatting into runs."""
    runs = []
    pos = 0
    while pos < len(text):
        bold_match = re.match(r"\*\*(.+?)\*\*", text[pos:])
        italic_match = re.match(r"\*(.+?)\*", text[pos:])
        if bold_match:
            runs.append({"text": bold_match.group(1), "bold": True, "italic": False,
                         "superscript": False, "subscript": False})
            pos += len(bold_match.group(0))
        elif italic_match:
            runs.append({"text": italic_match.group(1), "bold": False, "italic": True,
                         "superscript": False, "subscript": False})
            pos += len(italic_match.group(0))
        else:
            next_marker = len(text)
            for marker in ["**", "*"]:
                idx = text.find(marker, pos)
                if idx != -1 and idx < next_marker:
                    next_marker = idx
            chunk = text[pos:next_marker]
            if chunk:
                runs.append({"text": chunk, "bold": False, "italic": False,
                             "superscript": False, "subscript": False})
            pos = next_marker

    if not runs:
        runs = [{"text": text, "bold": False, "italic": False,
                 "superscript": False, "subscript": False}]
    return runs


# ---------------------------------------------------------------------------
# Pandoc JSON AST parser
# ---------------------------------------------------------------------------

def _pandoc_json_to_items(json_str: str, figures_path: Optional[Path] = None) -> list:
    """Convert Pandoc JSON AST to internal item format."""
    import json
    try:
        ast = json.loads(json_str)
        raw_items = []
        blocks = ast.get("blocks", [])
        for block in blocks:
            t = block.get("t", "")
            c = block.get("c", [])
            if t == "Header":
                level = c[0] if c else 1
                text = _pandoc_inlines_to_text(c[2] if len(c) > 2 else [])
                type_map = {1: "heading1", 2: "heading2", 3: "heading3"}
                raw_items.append({"type": type_map.get(level, "heading1"), "text": text,
                              "runs": [{"text": text, "bold": True, "italic": False,
                                        "superscript": False, "subscript": False}]})
            elif t == "Para":
                # Check if paragraph contains an image
                img_item = _pandoc_para_to_image(c, figures_path)
                if img_item:
                    raw_items.append(img_item)
                else:
                    text = _pandoc_inlines_to_text(c)
                    if text.strip():
                        raw_items.append({"type": "paragraph", "text": text,
                                      "runs": _pandoc_inlines_to_runs(c)})
            elif t == "Table":
                raw_items.append({"type": "table", "text": "", "runs": [], "rows": []})
        return _post_classify_items(raw_items)
    except Exception as e:
        log.warning("pandoc_json_parse_failed", error=str(e))
        return []


def _pandoc_para_to_image(inlines: list, figures_path: Optional[Path]) -> Optional[dict]:
    """If a Para block contains only an Image inline, return an image item."""
    # Filter out spaces
    non_space = [il for il in inlines if il.get("t") != "Space"]
    if len(non_space) == 1 and non_space[0].get("t") == "Image":
        img = non_space[0]
        c = img.get("c", [])
        # c = [attr, alt_inlines, [src, title]]
        alt_inlines = c[1] if len(c) > 1 else []
        src_info = c[2] if len(c) > 2 else ["", ""]
        src = src_info[0] if src_info else ""
        alt = _pandoc_inlines_to_text(alt_inlines)

        resolved = src
        if figures_path and src and not src.startswith(("http://", "https://", "data:")):
            candidate = figures_path / Path(src).name
            if candidate.exists():
                resolved = str(candidate)

        return {
            "type": "image",
            "path": resolved,
            "alt": alt or "Figure",
            "caption": alt or "",
        }
    return None


def _post_classify_items(items: list) -> list:
    """Post-process raw items to detect abstract, keywords, references, and supplementary."""
    result = []
    in_abstract = False
    in_references = False

    for item in items:
        itype = item["type"]

        # Pass image items through unchanged (no "text" key)
        if itype == "image":
            result.append(item)
            continue

        text = item.get("text", "").strip()

        # Detect section transitions via headings
        if itype in ("heading1", "heading2", "heading3"):
            tl = text.lower().rstrip(".")
            if tl == "abstract":
                item = dict(item, type="abstract_heading")
                in_abstract = True
                in_references = False
            elif tl in ("references", "bibliography", "works cited"):
                in_abstract = False
                in_references = True
            elif _SUPPLEMENTARY_RE.match(text):
                item = dict(item, type="supplementary_heading")
                in_abstract = False
                in_references = False
            else:
                in_abstract = False
            result.append(item)
            continue

        # Classify paragraphs based on context
        if in_abstract:
            if text.lower().startswith("keyword"):
                item = dict(item, type="keywords")
                in_abstract = False
            else:
                item = dict(item, type="abstract")
        elif in_references:
            if re.match(r"^\[?\d+\]?\.?\s+\w", text) and len(text) > 20:
                item = dict(item, type="reference")
            elif re.match(r"^\d+\.\s+\w", text) and len(text) > 20:
                item = dict(item, type="reference")
        elif not result and itype == "paragraph" and len(text) < 150 and not text.endswith("."):
            item = dict(item, type="title")

        result.append(item)

    return result


# ---------------------------------------------------------------------------
# Pandoc inline helpers
# ---------------------------------------------------------------------------

def _pandoc_inlines_to_text(inlines: list) -> str:
    parts = []
    for inline in inlines:
        t = inline.get("t", "")
        c = inline.get("c", "")
        if t == "Str":
            parts.append(c)
        elif t == "Space":
            parts.append(" ")
        elif t in ("Emph", "Strong"):
            parts.append(_pandoc_inlines_to_text(c))
    return "".join(parts)


def _pandoc_inlines_to_runs(inlines: list) -> list:
    runs = []
    for inline in inlines:
        t = inline.get("t", "")
        c = inline.get("c", "")
        if t == "Str":
            runs.append({"text": c, "bold": False, "italic": False,
                         "superscript": False, "subscript": False})
        elif t == "Space":
            runs.append({"text": " ", "bold": False, "italic": False,
                         "superscript": False, "subscript": False})
        elif t == "Strong":
            for r in _pandoc_inlines_to_runs(c):
                r["bold"] = True
                runs.append(r)
        elif t == "Emph":
            for r in _pandoc_inlines_to_runs(c):
                r["italic"] = True
                runs.append(r)
    return runs or [{"text": "", "bold": False, "italic": False,
                     "superscript": False, "subscript": False}]


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------

def _extract_runs(para) -> list:
    """Extract formatted runs from a python-docx paragraph."""
    runs = []
    for run in para.runs:
        if run.text:
            runs.append({
                "text": run.text,
                "bold": bool(run.bold),
                "italic": bool(run.italic),
                "superscript": bool(run.font.superscript) if run.font else False,
                "subscript": bool(run.font.subscript) if run.font else False,
            })
    if not runs and para.text:
        runs = [{"text": para.text, "bold": False, "italic": False,
                 "superscript": False, "subscript": False}]
    return runs


# ---------------------------------------------------------------------------
# Bibliography parsers
# ---------------------------------------------------------------------------

def _parse_ris(path: str) -> list:
    """Parse RIS bibliography file into list of records."""
    records = []
    current = {}
    with open(path, encoding="utf-8", errors="replace") as f:
        for line in f:
            line = line.rstrip()
            if line.startswith("ER  -"):
                if current:
                    records.append(current)
                    current = {}
            elif "  - " in line:
                tag, _, value = line.partition("  - ")
                tag = tag.strip()
                value = value.strip()
                if tag in current:
                    if isinstance(current[tag], list):
                        current[tag].append(value)
                    else:
                        current[tag] = [current[tag], value]
                else:
                    current[tag] = value
    return records


def _parse_bibtex(path: str) -> list:
    """Parse BibTeX file into list of records (simplified)."""
    records = []
    text = Path(path).read_text(encoding="utf-8", errors="replace")
    entries = re.findall(r"@\w+\{([^,]+),([^@]+)\}", text, re.DOTALL)
    for key, body in entries:
        record = {"TY": "JOUR", "ID": key.strip()}
        fields = re.findall(r"(\w+)\s*=\s*\{([^}]+)\}", body)
        field_map = {
            "author": "AU", "title": "TI", "journal": "JO",
            "year": "PY", "volume": "VL", "pages": "SP",
            "doi": "DO", "abstract": "AB", "booktitle": "BT",
        }
        for field, value in fields:
            ris_key = field_map.get(field.lower(), field.upper()[:2])
            record[ris_key] = value.strip()
        records.append(record)
    return records


# ---------------------------------------------------------------------------
# Metadata extraction
# ---------------------------------------------------------------------------

def _extract_metadata(items: list) -> dict:
    """Extract title, authors, abstract from items list."""
    metadata = {"title": "", "abstract": "", "keywords": ""}
    for item in items:
        if item["type"] == "title" and not metadata["title"]:
            metadata["title"] = item["text"]
        elif item["type"] == "abstract" and not metadata["abstract"]:
            metadata["abstract"] = item["text"]
        elif item["type"] == "keywords" and not metadata["keywords"]:
            metadata["keywords"] = item["text"]
    return metadata
