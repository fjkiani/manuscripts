"""
Shared image embedding helpers for all format plugins.
Handles DOCX picture insertion with graceful fallback.
"""

from pathlib import Path
from typing import Optional

import structlog

log = structlog.get_logger()

# Max width for embedded images in DOCX (in EMUs: 1 inch = 914400 EMU)
# 5.5 inches fits within standard margins for most styles
DEFAULT_WIDTH_INCHES = 5.5


def add_image_to_doc(doc, item: dict, width_inches: float = DEFAULT_WIDTH_INCHES):
    """
    Add an image item to a python-docx Document.

    Inserts the image centered, followed by a caption paragraph if present.
    Gracefully skips if the image file doesn't exist or can't be read.

    Returns the paragraph containing the image, or None on failure.
    """
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    path = item.get("path", "")
    caption = item.get("caption", "") or item.get("alt", "")

    if not path:
        return None

    img_path = Path(path)
    if not img_path.exists():
        log.warning("image_not_found", path=path)
        return None

    # Skip formats that python-docx / Pillow can't handle
    suffix = img_path.suffix.lower()
    if suffix in (".emf", ".wmf"):
        log.warning("image_format_skipped", path=path, reason="EMF/WMF not supported")
        return None

    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(width_inches))

        # Add caption if present
        if caption:
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap_p.add_run(caption)
            cap_run.italic = True
            cap_run.font.size = Pt(10)

        return p
    except Exception as e:
        log.warning("image_embed_failed", path=path, error=str(e))
        # Fallback: insert a placeholder paragraph
        p = doc.add_paragraph(f"[Figure: {img_path.name}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p


def image_to_latex(item: dict) -> str:
    """Return LaTeX figure environment for an image item."""
    path = item.get("path", "")
    caption = item.get("caption", "") or item.get("alt", "")
    alt = item.get("alt", "figure")

    if not path:
        return ""

    img_path = Path(path)
    # Use just the filename — Pandoc/LaTeX will resolve via --resource-path
    fname = img_path.name

    lines = [
        r"\begin{figure}[htbp]",
        r"  \centering",
        f"  \\includegraphics[width=\\linewidth]{{{fname}}}",
    ]
    if caption:
        lines.append(f"  \\caption{{{caption}}}")
    lines.append(f"  \\label{{fig:{alt.lower().replace(' ', '_')}}}")
    lines.append(r"\end{figure}")
    return "\n".join(lines)


def image_to_html(item: dict) -> str:
    """Return HTML figure element for an image item."""
    path = item.get("path", "")
    caption = item.get("caption", "") or item.get("alt", "")
    alt = item.get("alt", "Figure")

    if not path:
        return ""

    img_path = Path(path)
    # Use just the filename for HTML (served from same dir or data URI)
    fname = img_path.name

    html = f'<figure style="text-align:center;margin:1.5em 0;">\n'
    html += f'  <img src="{fname}" alt="{alt}" style="max-width:100%;height:auto;">\n'
    if caption:
        html += f'  <figcaption style="font-style:italic;font-size:0.9em;margin-top:0.5em;">{caption}</figcaption>\n'
    html += '</figure>'
    return html
