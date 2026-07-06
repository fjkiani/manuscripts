"""Elsevier journal format plugin — elsarticle style."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional
from app.citations.reference_engine import resolve_references

FORMAT_NAME = "Elsevier"
FORMAT_SUFFIX = "_elsevier"


def build(items: list, output_path: str, ris_data: Optional[list] = None, zotero_enabled: bool = False) -> str:
    items = resolve_references(items, ris_data, "elsevier")
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    section_counter = [0]

    for item in items:
        t = item["type"]
        text = item["text"]
        runs = item.get("runs", [])

        if t == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = "Times New Roman"
            p.paragraph_format.space_after = Pt(12)

        elif t == "abstract_heading":
            p = doc.add_paragraph()
            run = p.add_run("Abstract")
            run.bold = True
            run.font.size = Pt(12)

        elif t == "abstract":
            p = doc.add_paragraph()
            _add_runs(p, runs, Pt(12))
            p.paragraph_format.space_after = Pt(6)

        elif t == "keywords":
            p = doc.add_paragraph()
            run = p.add_run("Keywords: ")
            run.bold = True
            run.font.size = Pt(11)
            kw = text.replace("Keywords:", "").replace("keywords:", "").strip()
            p.add_run(kw).font.size = Pt(11)

        elif t == "heading1":
            section_counter[0] += 1
            p = doc.add_paragraph()
            run = p.add_run(f"{section_counter[0]}. {text}")
            run.bold = True
            run.font.size = Pt(13)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        elif t == "heading2":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)

        elif t == "heading3":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.italic = True
            run.font.size = Pt(12)

        elif t == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = Pt(18)  # 1.5x
            _add_runs(p, runs, Pt(12))

        elif t in ("table_caption", "figure_caption"):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)

        elif t == "reference":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.add_run(text).font.size = Pt(11)

        elif t == "table":
            rows = item.get("rows", [])
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                _apply_three_line_table(table)
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell.get("text", "")

    doc.save(output_path)
    return output_path


def _apply_three_line_table(table):
    """Apply three-line table style (top, header-bottom, table-bottom borders only)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    # Simplified: just use Table Grid style
    table.style = "Table Grid"


def _add_runs(para, runs: list, font_size=None):
    for r in runs:
        run = para.add_run(r.get("text", ""))
        run.bold = r.get("bold", False)
        run.italic = r.get("italic", False)
        run.font.name = "Times New Roman"
        if font_size:
            run.font.size = font_size
