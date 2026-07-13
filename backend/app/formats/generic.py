"""Generic publication-ready format plugin — clean academic style, no journal constraints."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional

from app.citations.reference_engine import resolve_references
from app.formats._image_helper import add_image_to_doc

FORMAT_NAME = "Generic"
FORMAT_SUFFIX = "_generic"


def build(items: list, output_path: str, ris_data: Optional[list] = None, zotero_enabled: bool = False) -> str:
    """Build a clean, publication-ready DOCX from items list."""
    # Resolve citations
    items = resolve_references(items, ris_data, "generic")

    doc = Document()

    # Page setup: A4, 1-inch margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    for item in items:
        _add_item(doc, item)

    doc.save(output_path)
    return output_path


def _add_item(doc: Document, item: dict):
    t = item["type"]
    text = item.get("text", "")
    runs = item.get("runs", [])

    if t == "title":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"

    elif t == "abstract_heading":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Abstract")
        run.bold = True
        run.font.size = Pt(12)

    elif t == "abstract":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.right_indent = Cm(1)
        _add_runs(p, runs, font_size=Pt(10), italic=True)

    elif t == "keywords":
        p = doc.add_paragraph()
        run = p.add_run("Keywords: ")
        run.bold = True
        run.font.size = Pt(10)
        kw_text = text.replace("Keywords:", "").replace("keywords:", "").strip()
        p.add_run(kw_text).font.size = Pt(10)

    elif t == "heading1":
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    elif t == "heading2":
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)

    elif t == "heading3":
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.size = Pt(11)

    elif t == "paragraph":
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.5)
        _add_runs(p, runs)

    elif t in ("table_caption", "figure_caption"):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif t == "reference":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.add_run(text).font.size = Pt(10)

    elif t == "image":
        add_image_to_doc(doc, item)

    elif t == "supplementary_heading":
        doc.add_page_break()
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    elif t == "table":
        rows = item.get("rows", [])
        if rows:
            table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    if j < len(table.rows[i].cells):
                        table.rows[i].cells[j].text = cell.get("text", "")


def _add_runs(para, runs: list, font_size=None, italic: bool = False):
    for run_data in runs:
        run = para.add_run(run_data.get("text", ""))
        run.bold = run_data.get("bold", False)
        run.italic = run_data.get("italic", False) or italic
        if font_size:
            run.font.size = font_size
        run.font.name = "Times New Roman"
