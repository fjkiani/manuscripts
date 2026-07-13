"""
CrisPRO Preprint format plugin.

Matches the layout of the CrisPRO bioRxiv submission PDF:
- Computer Modern serif (LaTeX default) / Times New Roman (DOCX)
- A4, 2.5 cm margins, single column
- Centered title block (title, authors, affiliations, date)
- Bold section headings, bold-italic sub-section headings
- Justified body text, 1.2× line spacing
- Booktabs-style tables (no vertical lines)
- Full-width figures with bold "Figure N:" caption label
- Numbered [1] references with hanging indent
- "For Research Use Only" footer (in LaTeX/PDF via renderer)
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Optional

from app.citations.reference_engine import resolve_references
from app.formats._image_helper import add_image_to_doc

FORMAT_NAME = "CrisPRO Preprint"
FORMAT_SUFFIX = "_crispro"

_BODY_FONT = "Times New Roman"
_BODY_SIZE = Pt(11)
_LINE_SPACING = 1.2  # multiple


def build(
    items: list,
    output_path: str,
    ris_data: Optional[list] = None,
    zotero_enabled: bool = False,
) -> str:
    """Build a CrisPRO-style preprint DOCX from items list."""
    items = resolve_references(items, ris_data, "generic")

    doc = Document()

    # Page setup: A4, 2.5 cm margins
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Default body style
    normal = doc.styles["Normal"]
    normal.font.name = _BODY_FONT
    normal.font.size = _BODY_SIZE
    _set_line_spacing(normal.paragraph_format, _LINE_SPACING)

    for item in items:
        _add_item(doc, item)

    doc.save(output_path)
    return output_path


def _add_item(doc: Document, item: dict):
    t = item["type"]
    text = item.get("text", "")
    runs = item.get("runs", [])

    if t == "title":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = _BODY_FONT
        run.font.size = Pt(20)
        # No bold — matches the PDF's normal-weight large title

    elif t in ("author", "affiliation"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = _BODY_FONT
        run.font.size = Pt(11)
        if t == "affiliation":
            run.italic = True

    elif t == "abstract_heading":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("Abstract")
        run.bold = True
        run.font.name = _BODY_FONT
        run.font.size = Pt(12)

    elif t == "abstract":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.right_indent = Cm(0)
        p.paragraph_format.space_after = Pt(6)
        _set_line_spacing(p.paragraph_format, _LINE_SPACING)
        _add_runs(p, runs, font_size=_BODY_SIZE)

    elif t == "keywords":
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        kw_label = p.add_run("Keywords: ")
        kw_label.bold = True
        kw_label.font.size = Pt(10)
        kw_label.font.name = _BODY_FONT
        kw_text = text.replace("Keywords:", "").replace("keywords:", "").strip()
        kw_run = p.add_run(kw_text)
        kw_run.font.size = Pt(10)
        kw_run.font.name = _BODY_FONT

    elif t == "heading1":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.name = _BODY_FONT
        run.font.size = Pt(12)

    elif t == "heading2":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.bold = True
        run.font.name = _BODY_FONT
        run.font.size = Pt(11)

    elif t == "heading3":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.italic = True
        run.font.name = _BODY_FONT
        run.font.size = Pt(11)

    elif t == "paragraph":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        _set_line_spacing(p.paragraph_format, _LINE_SPACING)
        _add_runs(p, runs)

    elif t in ("table_caption", "figure_caption"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        # Bold "Figure N:" label, then normal caption text
        label_end = text.find(":")
        if label_end != -1:
            label_run = p.add_run(text[: label_end + 1])
            label_run.bold = True
            label_run.font.size = Pt(10)
            label_run.font.name = _BODY_FONT
            rest_run = p.add_run(text[label_end + 1 :])
            rest_run.font.size = Pt(10)
            rest_run.font.name = _BODY_FONT
        else:
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = _BODY_FONT

    elif t == "reference":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = _BODY_FONT

    elif t == "image":
        add_image_to_doc(doc, item, width_inches=5.5)

    elif t == "supplementary_heading":
        doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.bold = True
        run.font.name = _BODY_FONT
        run.font.size = Pt(14)

    elif t == "table":
        rows = item.get("rows", [])
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=ncols)
        # Booktabs-style: no vertical borders, top/mid/bottom rules
        table.style = "Table Grid"
        _apply_booktabs_style(table)
        for i, row in enumerate(rows):
            for j, cell in enumerate(row):
                if j < len(table.rows[i].cells):
                    cell_obj = table.rows[i].cells[j]
                    cell_obj.text = cell.get("text", "")
                    if i == 0:
                        for run in cell_obj.paragraphs[0].runs:
                            run.bold = True


def _apply_booktabs_style(table):
    """Remove vertical borders from a table (booktabs style)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Remove table borders (we'll rely on row-level top/bottom rules)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("left", "right", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _set_line_spacing(para_format, multiple: float):
    """Set line spacing as a multiple."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    para_format.line_spacing = multiple


def _add_runs(para, runs: list, font_size=None, italic: bool = False):
    for run_data in runs:
        run = para.add_run(run_data.get("text", ""))
        run.bold = run_data.get("bold", False)
        run.italic = run_data.get("italic", False) or italic
        run.font.name = _BODY_FONT
        if font_size:
            run.font.size = font_size
        else:
            run.font.size = _BODY_SIZE
