"""Springer LNCS format plugin."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional
from app.citations.reference_engine import resolve_references
from app.formats._image_helper import add_image_to_doc

FORMAT_NAME = "Springer LNCS"
FORMAT_SUFFIX = "_springer"


def build(items: list, output_path: str, ris_data: Optional[list] = None, zotero_enabled: bool = False) -> str:
    items = resolve_references(items, ris_data, "springer")
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(15.2)
    section.page_height = Cm(23.5)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)

    sec_counter = [0]
    subsec_counter = [0]

    for item in items:
        t = item["type"]
        text = item.get("text", "")
        runs = item.get("runs", [])

        if t == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_after = Pt(12)

        elif t in ("abstract_heading",):
            p = doc.add_paragraph()
            run = p.add_run("Abstract.")
            run.bold = True
            run.font.size = Pt(10)

        elif t == "abstract":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.6)
            _add_runs(p, runs, Pt(10))

        elif t == "keywords":
            p = doc.add_paragraph()
            run = p.add_run("Keywords: ")
            run.bold = True
            run.font.size = Pt(10)
            kw = text.replace("Keywords:", "").replace("keywords:", "").strip()
            p.add_run(kw).font.size = Pt(10)

        elif t == "heading1":
            sec_counter[0] += 1
            subsec_counter[0] = 0
            p = doc.add_paragraph()
            run = p.add_run(f"{sec_counter[0]}  {text}")
            run.bold = True
            run.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(10)

        elif t == "heading2":
            subsec_counter[0] += 1
            p = doc.add_paragraph()
            run = p.add_run(f"{sec_counter[0]}.{subsec_counter[0]}  {text}")
            run.bold = True
            run.font.size = Pt(10)

        elif t == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.5)
            _add_runs(p, runs, Pt(10))

        elif t in ("table_caption", "figure_caption"):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(9)

        elif t == "reference":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.add_run(text).font.size = Pt(9)

        elif t == "image":
            add_image_to_doc(doc, item)

        elif t == "supplementary_heading":
            doc.add_page_break()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)

        elif t == "table":
            rows = item.get("rows", [])
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell.get("text", "")

    doc.save(output_path)
    return output_path


def _add_runs(para, runs: list, font_size=None):
    for r in runs:
        run = para.add_run(r.get("text", ""))
        run.bold = r.get("bold", False)
        run.italic = r.get("italic", False)
        run.font.name = "Times New Roman"
        if font_size:
            run.font.size = font_size
