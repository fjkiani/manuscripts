"""IEEE journal/conference format plugin — IEEEtran style."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Optional
from app.citations.reference_engine import resolve_references

FORMAT_NAME = "IEEE"
FORMAT_SUFFIX = "_ieee"


def build(items: list, output_path: str, ris_data: Optional[list] = None, zotero_enabled: bool = False) -> str:
    items = resolve_references(items, ris_data, "ieee")
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.59)   # Letter width
    section.page_height = Cm(27.94)  # Letter height
    section.left_margin = Cm(1.905)
    section.right_margin = Cm(1.905)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)

    ref_counter = [0]

    for item in items:
        t = item["type"]
        text = item["text"]
        runs = item.get("runs", [])

        if t == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(24)
            run.font.name = "Times New Roman"
            p.paragraph_format.space_after = Pt(6)

        elif t == "abstract_heading":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("Abstract")
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

        elif t == "abstract":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.right_indent = Cm(0.6)
            run = p.add_run(text)
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
            run.italic = True

        elif t == "keywords":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6)
            run = p.add_run("Index Terms—")
            run.bold = True
            run.font.size = Pt(9)
            kw = text.replace("Keywords:", "").replace("keywords:", "").strip()
            kw_run = p.add_run(kw)
            kw_run.font.size = Pt(9)

        elif t == "heading1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # IEEE uses Roman numerals for sections
            ref_counter[0] += 1
            run = p.add_run(f"{_to_roman(ref_counter[0])}. {text.upper()}")
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
            p.paragraph_format.space_before = Pt(6)

        elif t == "heading2":
            p = doc.add_paragraph()
            run = p.add_run(f"{text}")
            run.italic = True
            run.font.size = Pt(10)

        elif t == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.35)
            _add_runs(p, runs, Pt(10))

        elif t in ("table_caption", "figure_caption"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(8)

        elif t == "reference":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.add_run(text).font.size = Pt(8)

        elif t == "table":
            rows = item.get("rows", [])
            if rows:
                table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell.get("text", "")

    doc.save(output_path)
    return output_path


def _add_runs(para, runs: list, font_size=None):
    for r in runs:
        run = para.add_run(r.get("text", ""))
        run.bold = r.get("bold", False)
        run.italic = r.get("italic", False)
        run.font.name = "Times New Roman"
        if font_size:
            run.font.size = font_size


def _to_roman(n: int) -> str:
    vals = [(1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),(100,'C'),(90,'XC'),
            (50,'L'),(40,'XL'),(10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')]
    result = ""
    for v, s in vals:
        while n >= v:
            result += s
            n -= v
    return result
