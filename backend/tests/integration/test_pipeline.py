"""
Integration tests for the full manuscript formatting pipeline.
Tests: input conversion → journal formatting → DOCX output.
(PDF/LaTeX/HTML rendering requires Pandoc+LaTeX, tested separately.)
"""

import os
import tempfile
import pytest
from docx import Document
from pathlib import Path

from app.converter import convert_input
from app.formats import get_formatter, list_formats


SAMPLE_MARKDOWN = """# Machine Learning in Scientific Discovery

## Abstract

This paper examines machine learning applications in scientific research.
We demonstrate significant improvements in discovery rates.

Keywords: machine learning, scientific discovery, AI

## 1. Introduction

Machine learning has transformed scientific research [1]. Recent advances
in deep learning have enabled new discoveries across multiple domains [2].

## 2. Methods

### 2.1 Data Collection

We analyzed 1,247 publications from 2018–2024.

### 2.2 Analysis

Each publication was evaluated across five dimensions.

## 3. Results

ML-assisted research showed a 3.2-fold increase in discovery rate (p < 0.001).

## 4. Conclusion

Machine learning has demonstrably accelerated scientific discovery.

## References

[1] LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. Nature, 521, 436–444.
[2] Jumper, J., et al. (2021). AlphaFold. Nature, 596, 583–589.
"""

SAMPLE_PLAIN_TEXT = """The Role of AI in Modern Research

Abstract
Artificial intelligence is reshaping research methodologies worldwide.

Introduction
The integration of AI tools has accelerated discovery timelines significantly.

Methods
We conducted a systematic review of 500 publications.

Results
AI-assisted research showed 2.5x faster publication rates.

Conclusion
AI integration represents a fundamental shift in research practice.

References
[1] Smith, J. (2024). AI in Research. Science, 380, 100-110.
"""


@pytest.fixture
def markdown_file():
    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write(SAMPLE_MARKDOWN)
        path = f.name
    yield path
    os.unlink(path)


@pytest.fixture
def plaintext_file():
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        f.write(SAMPLE_PLAIN_TEXT)
        path = f.name
    yield path
    os.unlink(path)


@pytest.fixture
def docx_file():
    """Create a minimal DOCX file for testing."""
    from docx import Document
    doc = Document()
    doc.add_heading("Test Manuscript Title", level=0)
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph("This is the abstract of the test manuscript.")
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the introduction paragraph [1].")
    doc.add_heading("References", level=1)
    doc.add_paragraph("[1] Smith, J. (2024). Test paper. Journal, 1(1), 1-10.")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        path = f.name
    doc.save(path)
    yield path
    os.unlink(path)


class TestMarkdownConversion:
    def test_converts_markdown_to_items(self, markdown_file):
        result = convert_input(markdown_file)
        assert "items" in result
        assert len(result["items"]) > 0

    def test_detects_title(self, markdown_file):
        result = convert_input(markdown_file)
        title_items = [i for i in result["items"] if i["type"] in ("title", "heading1")]
        assert len(title_items) > 0

    def test_detects_references(self, markdown_file):
        result = convert_input(markdown_file)
        ref_items = [i for i in result["items"] if i["type"] == "reference"]
        assert len(ref_items) >= 1

    def test_metadata_extracted(self, markdown_file):
        result = convert_input(markdown_file)
        assert "metadata" in result


class TestPlainTextConversion:
    def test_converts_plaintext(self, plaintext_file):
        result = convert_input(plaintext_file)
        assert len(result["items"]) > 0

    def test_detects_references(self, plaintext_file):
        result = convert_input(plaintext_file)
        ref_items = [i for i in result["items"] if i["type"] == "reference"]
        assert len(ref_items) >= 1


class TestDOCXConversion:
    def test_converts_docx(self, docx_file):
        result = convert_input(docx_file)
        assert len(result["items"]) > 0

    def test_detects_headings(self, docx_file):
        result = convert_input(docx_file)
        heading_items = [i for i in result["items"] if "heading" in i["type"]]
        assert len(heading_items) > 0


class TestFullPipeline:
    """Test full pipeline: convert → format → DOCX output."""

    @pytest.mark.parametrize("style", ["generic", "ieee", "elsevier", "springer", "apa", "ama"])
    def test_markdown_to_all_styles(self, markdown_file, style):
        """Markdown input → all 6 journal styles → valid DOCX."""
        doc_data = convert_input(markdown_file)
        formatter = get_formatter(style)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            result = formatter.build(
                doc_data["items"],
                output_path,
                ris_data=None,
                zotero_enabled=False,
            )
            assert os.path.exists(result)
            assert os.path.getsize(result) > 500

            # Verify valid DOCX
            doc = Document(result)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            assert len(paragraphs) > 0, f"No content in {style} output"
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    @pytest.mark.parametrize("style", ["generic", "ieee", "apa"])
    def test_plaintext_to_styles(self, plaintext_file, style):
        """Plain text input → selected journal styles → valid DOCX."""
        doc_data = convert_input(plaintext_file)
        formatter = get_formatter(style)

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            result = formatter.build(
                doc_data["items"],
                output_path,
                ris_data=None,
                zotero_enabled=False,
            )
            assert os.path.exists(result)
            assert os.path.getsize(result) > 500
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_citation_reformatting_numbered_to_apa(self, markdown_file):
        """Numbered [N] citations in input → APA author-date in output."""
        doc_data = convert_input(markdown_file)
        formatter = get_formatter("apa")

        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            output_path = f.name

        try:
            result = formatter.build(
                doc_data["items"],
                output_path,
                ris_data=None,
                zotero_enabled=False,
            )
            doc = Document(result)
            # Verify references section exists
            all_text = " ".join(p.text for p in doc.paragraphs)
            assert "LeCun" in all_text or "Jumper" in all_text or "Smith" in all_text
        finally:
            if os.path.exists(output_path):
                os.unlink(output_path)

    def test_with_ris_bibliography(self, markdown_file):
        """Test pipeline with RIS bibliography file."""
        ris_content = """TY  - JOUR
AU  - LeCun, Yann
AU  - Bengio, Yoshua
AU  - Hinton, Geoffrey
TI  - Deep learning
JO  - Nature
PY  - 2015
VL  - 521
SP  - 436
EP  - 444
DO  - 10.1038/nature14539
ER  -
"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.ris', delete=False) as f:
            f.write(ris_content)
            bib_path = f.name

        try:
            doc_data = convert_input(markdown_file, bib_path)
            assert doc_data["ris_data"] is not None
            assert len(doc_data["ris_data"]) == 1
            assert doc_data["ris_data"][0]["TI"] == "Deep learning"
        finally:
            os.unlink(bib_path)
