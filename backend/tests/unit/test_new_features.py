"""Tests for image embedding, supplementary sections, and new format styles."""

import os
import tempfile
import pytest
from pathlib import Path
from docx import Document


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_png(path: Path) -> None:
    """Write a minimal valid 1×1 white PNG to *path*."""
    import struct, zlib

    def _chunk(tag: bytes, data: bytes) -> bytes:
        c = struct.pack(">I", len(data)) + tag + data
        return c + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\xFF\xFF\xFF")
    png = (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", ihdr)
        + _chunk(b"IDAT", idat)
        + _chunk(b"IEND", b"")
    )
    path.write_bytes(png)


def _sample_items(png_path: str = "") -> list:
    """Return a minimal item list that includes image + supplementary_heading."""
    return [
        {"type": "title", "text": "Test Title", "runs": [
            {"text": "Test Title", "bold": True, "italic": False, "superscript": False, "subscript": False}
        ]},
        {"type": "abstract_heading", "text": "Abstract", "runs": [
            {"text": "Abstract", "bold": True, "italic": False, "superscript": False, "subscript": False}
        ]},
        {"type": "abstract", "text": "This is the abstract.", "runs": [
            {"text": "This is the abstract.", "bold": False, "italic": False, "superscript": False, "subscript": False}
        ]},
        {"type": "heading1", "text": "Introduction", "runs": [
            {"text": "Introduction", "bold": True, "italic": False, "superscript": False, "subscript": False}
        ]},
        {"type": "paragraph", "text": "Body text.", "runs": [
            {"text": "Body text.", "bold": False, "italic": False, "superscript": False, "subscript": False}
        ]},
        {"type": "image", "path": png_path, "alt": "Figure 1", "caption": "A test figure."},
        {"type": "supplementary_heading", "text": "Supplementary Figures", "runs": [
            {"text": "Supplementary Figures", "bold": True, "italic": False, "superscript": False, "subscript": False}
        ]},
        {"type": "paragraph", "text": "Supplementary content.", "runs": [
            {"text": "Supplementary content.", "bold": False, "italic": False, "superscript": False, "subscript": False}
        ]},
        {"type": "reference_heading", "text": "References", "runs": [
            {"text": "References", "bold": True, "italic": False, "superscript": False, "subscript": False}
        ]},
        {"type": "reference", "text": "[1] Author et al. 2024.", "runs": [
            {"text": "[1] Author et al. 2024.", "bold": False, "italic": False, "superscript": False, "subscript": False}
        ]},
    ]


# ---------------------------------------------------------------------------
# Converter: supplementary_heading detection
# ---------------------------------------------------------------------------

class TestSupplementaryDetection:
    def test_supplementary_heading_detected(self):
        from app.converter import _post_classify_items
        items = [{"type": "heading1", "text": "Supplementary Figures", "runs": []}]
        result = _post_classify_items(items)
        assert result[0]["type"] == "supplementary_heading"

    def test_supplemental_heading_detected(self):
        from app.converter import _post_classify_items
        items = [{"type": "heading1", "text": "Supplemental Data", "runs": []}]
        result = _post_classify_items(items)
        assert result[0]["type"] == "supplementary_heading"

    def test_appendix_heading_detected(self):
        from app.converter import _post_classify_items
        items = [{"type": "heading2", "text": "Appendix A", "runs": []}]
        result = _post_classify_items(items)
        assert result[0]["type"] == "supplementary_heading"

    def test_normal_heading_not_reclassified(self):
        from app.converter import _post_classify_items
        items = [{"type": "heading1", "text": "Introduction", "runs": []}]
        result = _post_classify_items(items)
        assert result[0]["type"] == "heading1"

    def test_image_items_pass_through(self):
        from app.converter import _post_classify_items
        items = [{"type": "image", "path": "/tmp/fig.png", "alt": "Figure 1", "caption": ""}]
        result = _post_classify_items(items)
        assert result[0]["type"] == "image"


# ---------------------------------------------------------------------------
# Format plugins: image + supplementary_heading via build()
# ---------------------------------------------------------------------------

# Legacy plugins expose build(items, output_path) at module level.
# New plugins (generic, crispro, preprint) also expose build() + _add_item().
LEGACY_STYLES = ["ieee", "elsevier", "springer", "apa", "ama"]
ALL_STYLES = ["generic", "ieee", "elsevier", "springer", "apa", "ama", "crispro", "preprint"]


@pytest.mark.parametrize("style_id", ALL_STYLES)
class TestImageAndSupplementaryViaBuilder:
    def test_build_with_image_and_supplementary(self, style_id, tmp_path):
        """build() must not raise when given image + supplementary_heading items."""
        import importlib
        mod = importlib.import_module(f"app.formats.{style_id}")

        png = tmp_path / "fig_001.png"
        _make_png(png)
        items = _sample_items(str(png))
        out = str(tmp_path / f"out_{style_id}.docx")

        mod.build(items, out)  # must not raise
        assert Path(out).exists() and Path(out).stat().st_size > 0

    def test_build_with_missing_image_no_crash(self, style_id, tmp_path):
        """build() must not raise when image path does not exist."""
        import importlib
        mod = importlib.import_module(f"app.formats.{style_id}")

        items = [
            {"type": "title", "text": "Test", "runs": [
                {"text": "Test", "bold": True, "italic": False, "superscript": False, "subscript": False}
            ]},
            {"type": "image", "path": str(tmp_path / "nonexistent.png"), "alt": "Missing", "caption": ""},
        ]
        out = str(tmp_path / f"out_missing_{style_id}.docx")
        mod.build(items, out)  # must not raise


# ---------------------------------------------------------------------------
# _add_item API for generic/crispro (they expose it separately)
# ---------------------------------------------------------------------------

@pytest.mark.parametrize("style_id", ["generic", "crispro"])
class TestAddItemAPI:
    def test_image_item_no_crash(self, style_id, tmp_path):
        import importlib
        mod = importlib.import_module(f"app.formats.{style_id}")
        png = tmp_path / "fig.png"
        _make_png(png)
        doc = Document()
        mod._add_item(doc, {"type": "image", "path": str(png), "alt": "Fig", "caption": ""})

    def test_supplementary_heading_adds_paragraph(self, style_id, tmp_path):
        import importlib
        mod = importlib.import_module(f"app.formats.{style_id}")
        doc = Document()
        before = len(doc.paragraphs)
        mod._add_item(doc, {"type": "supplementary_heading", "text": "Supplementary Figures", "runs": [
            {"text": "Supplementary Figures", "bold": True, "italic": False, "superscript": False, "subscript": False}
        ]})
        assert len(doc.paragraphs) > before


# ---------------------------------------------------------------------------
# New styles: crispro + preprint
# ---------------------------------------------------------------------------

class TestCrisPROFormatter:
    def test_crispro_builds_docx(self, tmp_path):
        from app.formats import crispro
        out = str(tmp_path / "crispro_test.docx")
        crispro.build(_sample_items(), out)
        assert Path(out).exists() and Path(out).stat().st_size > 0

    def test_crispro_page_size_a4(self, tmp_path):
        from app.formats import crispro
        out = str(tmp_path / "crispro_a4.docx")
        crispro.build(_sample_items(), out)
        doc = Document(out)
        section = doc.sections[0]
        # A4: 21 cm wide, 29.7 cm tall (within 1 mm tolerance)
        assert abs(section.page_width.cm - 21.0) < 0.2
        assert abs(section.page_height.cm - 29.7) < 0.2


class TestPreprintFormatter:
    def test_preprint_builds_docx(self, tmp_path):
        from app.formats import preprint
        out = str(tmp_path / "preprint_test.docx")
        preprint.build(_sample_items(), out)
        assert Path(out).exists() and Path(out).stat().st_size > 0

    def test_preprint_page_size_a4(self, tmp_path):
        from app.formats import preprint
        out = str(tmp_path / "preprint_a4.docx")
        preprint.build(_sample_items(), out)
        doc = Document(out)
        section = doc.sections[0]
        assert abs(section.page_width.cm - 21.0) < 0.2
        assert abs(section.page_height.cm - 29.7) < 0.2


# ---------------------------------------------------------------------------
# Formats registry
# ---------------------------------------------------------------------------

class TestFormatsRegistry:
    def test_all_eight_styles_available(self):
        from app.formats import list_formats
        ids = {f["id"] for f in list_formats()}
        expected = {"ieee", "elsevier", "springer", "apa", "ama", "generic", "crispro", "preprint"}
        assert expected == ids

    def test_crispro_in_registry(self):
        from app.formats import list_formats
        ids = {f["id"] for f in list_formats()}
        assert "crispro" in ids

    def test_preprint_in_registry(self):
        from app.formats import list_formats
        ids = {f["id"] for f in list_formats()}
        assert "preprint" in ids

    def test_unknown_style_raises(self):
        from app.formats import get_formatter
        with pytest.raises(Exception):
            get_formatter("nonexistent_style_xyz")
