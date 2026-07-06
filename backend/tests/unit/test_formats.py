"""Unit tests for journal format plugins."""

import os
import tempfile
import pytest
from docx import Document


SAMPLE_ITEMS = [
    {"type": "title", "text": "A Test Manuscript Title", "runs": [
        {"text": "A Test Manuscript Title", "bold": True, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "abstract_heading", "text": "Abstract", "runs": [
        {"text": "Abstract", "bold": True, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "abstract", "text": "This is the abstract of the test manuscript.", "runs": [
        {"text": "This is the abstract of the test manuscript.", "bold": False, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "keywords", "text": "Keywords: testing, manuscript, formatting", "runs": [
        {"text": "Keywords: testing, manuscript, formatting", "bold": False, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "heading1", "text": "Introduction", "runs": [
        {"text": "Introduction", "bold": True, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "paragraph", "text": "This is the introduction paragraph with some text [1].", "runs": [
        {"text": "This is the introduction paragraph with some text [1].", "bold": False, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "heading2", "text": "Background", "runs": [
        {"text": "Background", "bold": True, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "paragraph", "text": "Background information goes here.", "runs": [
        {"text": "Background information goes here.", "bold": False, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "heading1", "text": "Methods", "runs": [
        {"text": "Methods", "bold": True, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "paragraph", "text": "We used the following methods.", "runs": [
        {"text": "We used the following methods.", "bold": False, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "table_caption", "text": "Table 1. Summary of results.", "runs": [
        {"text": "Table 1. Summary of results.", "bold": True, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "table", "text": "", "runs": [], "rows": [
        [{"text": "Column A", "gridspan": 1, "runs": [{"text": "Column A", "bold": True, "italic": False, "superscript": False, "subscript": False}], "vmerge_continue": False},
         {"text": "Column B", "gridspan": 1, "runs": [{"text": "Column B", "bold": True, "italic": False, "superscript": False, "subscript": False}], "vmerge_continue": False}],
        [{"text": "Value 1", "gridspan": 1, "runs": [{"text": "Value 1", "bold": False, "italic": False, "superscript": False, "subscript": False}], "vmerge_continue": False},
         {"text": "Value 2", "gridspan": 1, "runs": [{"text": "Value 2", "bold": False, "italic": False, "superscript": False, "subscript": False}], "vmerge_continue": False}],
    ]},
    {"type": "heading1", "text": "Conclusion", "runs": [
        {"text": "Conclusion", "bold": True, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "paragraph", "text": "In conclusion, this test demonstrates the formatting pipeline.", "runs": [
        {"text": "In conclusion, this test demonstrates the formatting pipeline.", "bold": False, "italic": False, "superscript": False, "subscript": False}
    ]},
    {"type": "reference", "text": "[1] Smith, J. (2024). Test paper. Journal, 1(1), 1-10.", "runs": [
        {"text": "[1] Smith, J. (2024). Test paper. Journal, 1(1), 1-10.", "bold": False, "italic": False, "superscript": False, "subscript": False}
    ]},
]


def _build_and_verify(module_name: str) -> str:
    """Helper: build a DOCX with the given format plugin and verify it's valid."""
    import importlib
    mod = importlib.import_module(f"app.formats.{module_name}")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        output_path = f.name

    try:
        result = mod.build(SAMPLE_ITEMS.copy(), output_path, ris_data=None, zotero_enabled=False)
        assert os.path.exists(result), f"Output file not created: {result}"
        assert os.path.getsize(result) > 1000, f"Output file too small: {os.path.getsize(result)} bytes"

        # Verify it's a valid DOCX
        doc = Document(result)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(paragraphs) > 0, "No paragraphs in output document"

        return result
    finally:
        if os.path.exists(output_path) and output_path != result:
            os.unlink(output_path)


class TestGenericFormat:
    def test_builds_valid_docx(self):
        path = _build_and_verify("generic")
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("Test Manuscript" in t for t in texts)
        os.unlink(path)

    def test_has_title(self):
        path = _build_and_verify("generic")
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("A Test Manuscript Title" in t for t in texts)
        os.unlink(path)

    def test_has_references(self):
        path = _build_and_verify("generic")
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("Smith" in t for t in texts)
        os.unlink(path)


class TestIEEEFormat:
    def test_builds_valid_docx(self):
        path = _build_and_verify("ieee")
        assert os.path.getsize(path) > 1000
        os.unlink(path)

    def test_has_content(self):
        path = _build_and_verify("ieee")
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(texts) > 3
        os.unlink(path)


class TestElsevierFormat:
    def test_builds_valid_docx(self):
        path = _build_and_verify("elsevier")
        assert os.path.getsize(path) > 1000
        os.unlink(path)


class TestSpringerFormat:
    def test_builds_valid_docx(self):
        path = _build_and_verify("springer")
        assert os.path.getsize(path) > 1000
        os.unlink(path)


class TestAPAFormat:
    def test_builds_valid_docx(self):
        path = _build_and_verify("apa")
        assert os.path.getsize(path) > 1000
        os.unlink(path)


class TestAMAFormat:
    def test_builds_valid_docx(self):
        path = _build_and_verify("ama")
        assert os.path.getsize(path) > 1000
        os.unlink(path)


class TestFormatRegistry:
    def test_all_formats_registered(self):
        from app.formats import list_formats
        formats = list_formats()
        format_ids = {f["id"] for f in formats}
        expected = {"ieee", "elsevier", "springer", "apa", "ama", "generic"}
        assert expected.issubset(format_ids), f"Missing formats: {expected - format_ids}"

    def test_get_formatter(self):
        from app.formats import get_formatter
        for style in ["ieee", "elsevier", "springer", "apa", "ama", "generic"]:
            formatter = get_formatter(style)
            assert hasattr(formatter, "build")
            assert hasattr(formatter, "FORMAT_NAME")

    def test_invalid_style_raises(self):
        from app.formats import get_formatter
        with pytest.raises(ValueError):
            get_formatter("nonexistent_style")
